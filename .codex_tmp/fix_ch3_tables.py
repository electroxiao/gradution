from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = Path(r"C:\Users\Clark\Desktop\毕业设计\第3章草稿.docx")
OUT = Path(r"C:\dev\gradution\.codex_tmp\第3章草稿_edited.docx")


def iter_blocks(doc):
    p_i = t_i = 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield ("p", p_i, doc.paragraphs[p_i], child)
            p_i += 1
        elif child.tag == qn("w:tbl"):
            yield ("tbl", t_i, doc.tables[t_i], child)
            t_i += 1


def ensure_before_half_line(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), "157")
    spacing.set(qn("w:beforeLines"), "50")


def is_table_caption(paragraph):
    return paragraph.text.strip().startswith("表3-")


doc = Document(SRC)
blocks = list(iter_blocks(doc))

caption_style = None
for kind, _, obj, _ in blocks:
    if kind == "p" and is_table_caption(obj):
        caption_style = obj.style
        break

start = None
end = None
for i, (kind, _, obj, _) in enumerate(blocks):
    if kind == "p" and obj.style.name.startswith("Heading 2") and "业务模块设计" in obj.text:
        start = i
        continue
    if start is not None and i > start and kind == "p" and obj.style.name.startswith("Heading 2"):
        end = i
        break

if start is None:
    raise RuntimeError("未找到“业务模块设计”表格区。")
if end is None:
    end = len(blocks)

changed_after = []
normalized_captions = []
missing_captions = []

for i in range(start, end):
    kind, table_idx, _, _ = blocks[i]
    if kind != "tbl":
        continue

    prev_block = blocks[i - 1] if i > 0 else None
    if prev_block is None or prev_block[0] != "p" or not is_table_caption(prev_block[2]):
        missing_captions.append(table_idx)
    elif caption_style is not None and prev_block[2].style != caption_style:
        prev_block[2].style = caption_style
        normalized_captions.append(table_idx)

    next_block = blocks[i + 1] if i + 1 < len(blocks) else None
    if next_block is not None and next_block[0] == "p":
        ensure_before_half_line(next_block[2])
        changed_after.append(table_idx)

doc.save(OUT)

print(f"saved={OUT}")
print(f"section_blocks={start}-{end}")
print(f"tables_after_spacing={changed_after}")
print(f"normalized_captions={normalized_captions}")
print(f"missing_captions={missing_captions}")
