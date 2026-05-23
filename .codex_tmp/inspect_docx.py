from pathlib import Path
import json
from docx import Document
from docx.oxml.ns import qn

DOCX = Path(r"C:\Users\Clark\Desktop\毕业设计\第3章草稿.docx")


def iter_blocks(doc):
    body = doc.element.body
    p_i = t_i = 0
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            p = doc.paragraphs[p_i]
            yield ("p", p_i, p)
            p_i += 1
        elif child.tag == qn("w:tbl"):
            t = doc.tables[t_i]
            yield ("tbl", t_i, t)
            t_i += 1


def text_of_table(table):
    rows = []
    for row in table.rows[:3]:
        rows.append(" | ".join(cell.text.replace("\n", " / ").strip() for cell in row.cells))
    return " || ".join(rows)


blocks = list(iter_blocks(Document(DOCX)))
start = None
end = None
for i, (kind, idx, obj) in enumerate(blocks):
    if kind == "p":
        text = obj.text.strip()
        if "3.4.3" in text and start is None:
            start = i
        elif start is not None and i > start and (
            text.startswith("3.4.4")
            or text.startswith("3.5")
            or text.startswith("第四章")
            or text.startswith("4.")
        ):
            end = i
            break

print(f"blocks={len(blocks)} tables={len(Document(DOCX).tables)} start={start} end={end}")

lo = max(0, (start or 0) - 25)
hi = min(len(blocks), (end or len(blocks)) + 8)
for n in range(lo, hi):
    kind, idx, obj = blocks[n]
    if kind == "p":
        p = obj
        text = p.text.strip().replace("\n", " / ")
        if len(text) > 120:
            text = text[:117] + "..."
        pf = p.paragraph_format
        before = pf.space_before
        payload = {
            "block": n,
            "type": "P",
            "idx": idx,
            "style": p.style.name,
            "before_pt": before.pt if before else None,
            "align": str(p.alignment),
            "text": text,
        }
        print(json.dumps(payload, ensure_ascii=True))
    else:
        payload = {
            "block": n,
            "type": "T",
            "idx": idx,
            "rows": len(obj.rows),
            "cols": len(obj.columns),
            "text": text_of_table(obj),
        }
        print(json.dumps(payload, ensure_ascii=True))
