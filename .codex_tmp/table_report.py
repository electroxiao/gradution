from pathlib import Path
import json
from docx import Document
from docx.oxml.ns import qn

DOCX = Path(r"C:\Users\Clark\Desktop\毕业设计\第3章草稿.docx")


def iter_blocks(doc):
    p_i = t_i = 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield ("p", p_i, doc.paragraphs[p_i])
            p_i += 1
        elif child.tag == qn("w:tbl"):
            yield ("tbl", t_i, doc.tables[t_i])
            t_i += 1


def text_table(tbl):
    if not tbl.rows:
        return ""
    return " | ".join(cell.text.replace("\n", " / ").strip() for cell in tbl.rows[0].cells)


blocks = list(iter_blocks(Document(DOCX)))
for n, (kind, idx, obj) in enumerate(blocks):
    if kind != "tbl":
        continue
    prev = blocks[n - 1] if n else None
    nxt = blocks[n + 1] if n + 1 < len(blocks) else None
    nxt2 = blocks[n + 2] if n + 2 < len(blocks) else None
    rec = {
        "block": n,
        "table": idx,
        "header": text_table(obj),
        "prev_kind": prev[0] if prev else None,
        "prev_style": prev[2].style.name if prev and prev[0] == "p" else None,
        "prev_text": prev[2].text if prev and prev[0] == "p" else None,
        "next_kind": nxt[0] if nxt else None,
        "next_style": nxt[2].style.name if nxt and nxt[0] == "p" else None,
        "next_before_pt": (nxt[2].paragraph_format.space_before.pt if nxt and nxt[0] == "p" and nxt[2].paragraph_format.space_before else None),
        "next_text": nxt[2].text if nxt and nxt[0] == "p" else None,
        "next2_kind": nxt2[0] if nxt2 else None,
        "next2_style": nxt2[2].style.name if nxt2 and nxt2[0] == "p" else None,
        "next2_before_pt": (nxt2[2].paragraph_format.space_before.pt if nxt2 and nxt2[0] == "p" and nxt2[2].paragraph_format.space_before else None),
        "next2_text": nxt2[2].text if nxt2 and nxt2[0] == "p" else None,
    }
    print(json.dumps(rec, ensure_ascii=True))
