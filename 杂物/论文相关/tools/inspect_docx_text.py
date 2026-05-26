from pathlib import Path

from docx import Document


BASE = Path(r"C:\Users\Clark\Desktop\毕业设计")
FILES = [
    "新第3章草稿_3.4.3表题已处理.docx",
    "新第5章草稿.docx",
]


for name in FILES:
    path = BASE / name
    doc = Document(str(path))
    texts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            texts.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        texts.append("[TABLE] " + text)

    print(f"--- {name} texts={len(texts)}")
    for text in texts[:100]:
        print(text)
