from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "毕业论文" / "第4章智能助学系统设计初稿.md"
OUTPUT = ROOT / "毕业论文" / "第4章智能助学系统设计.docx"


def set_run_font(run, east_asia: str = "宋体", ascii_font: str = "Times New Roman", size: int | None = None) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size:
        run.font.size = Pt(size)


def set_paragraph_text(paragraph, text: str, bold: bool = False, size: int | None = None) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    set_run_font(run, size=size)


def set_cell_text(cell, text: str) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_text(paragraph, clean_inline(text), size=10)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def clean_inline(text: str) -> str:
    text = text.strip()
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    return text


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def is_separator_line(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        if is_separator_line(line):
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def add_table(document: Document, lines: list[str]) -> None:
    rows = parse_table(lines)
    if not rows:
        return

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.autofit = True

    for row_index, row in enumerate(rows):
        for col_index in range(column_count):
            text = row[col_index] if col_index < len(row) else ""
            cell = table.cell(row_index, col_index)
            set_cell_text(cell, text)
            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    document.add_paragraph()


def add_placeholder(document: Document, caption: str | None) -> None:
    label = caption or "图占位符"
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_text(paragraph, f"[{label}占位符：请在此处插入对应图片]", bold=True)


def build_docx() -> None:
    markdown = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()

    section = document.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(11)

    pending_figure_caption: str | None = None
    in_mermaid = False
    in_code = False
    table_buffer: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_table(document, table_buffer)
            table_buffer = []

    for raw_line in markdown:
        line = raw_line.rstrip()

        if in_mermaid:
            if line.strip() == "```":
                add_placeholder(document, pending_figure_caption)
                pending_figure_caption = None
                in_mermaid = False
            continue

        if in_code:
            if line.strip() == "```":
                in_code = False
            else:
                paragraph = document.add_paragraph(style=None)
                run = paragraph.add_run(line)
                run.font.name = "Courier New"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run.font.size = Pt(9)
            continue

        if line.strip() == "```mermaid":
            flush_table()
            in_mermaid = True
            continue

        if line.strip().startswith("```"):
            flush_table()
            in_code = True
            continue

        if is_table_line(line):
            table_buffer.append(line)
            continue
        flush_table()

        if not line.strip():
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            text = clean_inline(heading_match.group(2))
            paragraph = document.add_heading(text, level=level)
            for run in paragraph.runs:
                set_run_font(run, east_asia="黑体", size=16 if level == 1 else 13)
            continue

        caption_match = re.match(r"^\*\*(图4-\d+|表4-\d+)\s+(.+)\*\*$", line.strip())
        if caption_match:
            label = f"{caption_match.group(1)} {caption_match.group(2)}"
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_text(paragraph, label, bold=True)
            if label.startswith("图"):
                pending_figure_caption = label
            continue

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(22)
        paragraph.paragraph_format.line_spacing = 1.5
        set_paragraph_text(paragraph, clean_inline(line))

    flush_table()
    document.save(OUTPUT)


if __name__ == "__main__":
    build_docx()
    print(OUTPUT)
